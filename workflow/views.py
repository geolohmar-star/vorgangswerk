# SPDX-License-Identifier: EUPL-1.2
# Copyright (C) 2026 Georg Klein
"""Workflow-App Views: Arbeitsstapel, Task-Verwaltung, Editor, Prozesszentrale."""
import json
import logging

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.contenttypes.models import ContentType
from django.db.models import Count, Q
from django.http import JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.http import require_POST

from .forms import ProzessAntragForm, WorkflowTriggerForm
from .models import (
    ProzessAntrag,
    WorkflowInstance,
    WorkflowStep,
    WorkflowTask,
    WorkflowTemplate,
    WorkflowTransition,
    WorkflowTrigger,
)
from .services import WorkflowEngine

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Arbeitsstapel
# ---------------------------------------------------------------------------


@login_required
def arbeitsstapel(request):
    """Zeigt alle offenen Tasks des eingeloggten Nutzers.

    Enthaelt sowohl direkt zugewiesene als auch Gruppen-Tasks.
    """
    user = request.user
    user_gruppen = user.groups.all()

    # Direkt zugewiesene Tasks
    eigene_tasks = WorkflowTask.objects.filter(
        zugewiesen_an_user=user,
        status__in=[WorkflowTask.STATUS_OFFEN, WorkflowTask.STATUS_IN_BEARBEITUNG],
    ).select_related("step", "step__template", "instance")

    # Gruppen-Tasks (nicht geclaimed)
    gruppen_tasks = WorkflowTask.objects.filter(
        zugewiesen_an_gruppe__in=user_gruppen,
        status=WorkflowTask.STATUS_OFFEN,
        claimed_von__isnull=True,
    ).select_related("step", "step__template", "instance", "zugewiesen_an_gruppe")

    # Vom User geclaimte Gruppen-Tasks
    geclaimte_tasks = WorkflowTask.objects.filter(
        claimed_von=user,
        status__in=[WorkflowTask.STATUS_OFFEN, WorkflowTask.STATUS_IN_BEARBEITUNG],
    ).exclude(zugewiesen_an_user=user).select_related(
        "step", "step__template", "instance"
    )

    kontext = {
        "eigene_tasks": eigene_tasks,
        "gruppen_tasks": gruppen_tasks,
        "geclaimte_tasks": geclaimte_tasks,
        "eigene_anzahl": eigene_tasks.count(),
        "gruppen_anzahl": gruppen_tasks.count(),
        "geclaimte_anzahl": geclaimte_tasks.count(),
    }
    return render(request, "workflow/arbeitsstapel.html", kontext)


@login_required
def task_detail(request, pk):
    """Zeigt Detail-Ansicht eines Tasks und erlaubt Abschluss."""
    task = get_object_or_404(
        WorkflowTask.objects.select_related(
            "step",
            "step__template",
            "instance",
            "zugewiesen_an_gruppe",
            "zugewiesen_an_user",
            "instance__aktueller_schritt",
        ),
        pk=pk,
    )

    if not task.kann_bearbeiten(request.user) and not request.user.is_staff:
        messages.error(request, "Du bist nicht berechtigt, diesen Task zu bearbeiten.")
        return redirect("workflow:arbeitsstapel")

    content_object = task.instance.content_object

    if request.method == "POST":
        entscheidung = request.POST.get("entscheidung", "genehmigt")
        kommentar = request.POST.get("kommentar", "")
        engine = WorkflowEngine()
        engine.complete_task(task, request.user, entscheidung=entscheidung, kommentar=kommentar)
        from core.models import audit
        audit(
            request,
            aktion="geaendert",
            app="workflow",
            objekt_typ="WorkflowTask",
            objekt_id=task.pk,
            beschreibung=(
                f"Task '{task.step.titel}' abgeschlossen – "
                f"Workflow: {task.step.template.name} – "
                f"Entscheidung: {entscheidung}"
                + (f" – Kommentar: {kommentar[:100]}" if kommentar else "")
            ),
        )
        messages.success(request, f"Task '{task.step.titel}' wurde abgeschlossen.")
        return redirect("workflow:arbeitsstapel")

    # Naechste moegliche Entscheidungen aus den Transitions lesen
    ausgaenge = WorkflowTransition.objects.filter(
        von_schritt=task.step,
        bedingung_typ="entscheidung",
    ).values_list("bedingung_entscheidung", flat=True)
    hat_ausgaenge = ausgaenge.exists()

    # Antragsdaten aufbereiten wenn content_object eine AntrSitzung ist
    antrag_felder = []
    sitzung = None
    try:
        from formulare.models import AntrSitzung
        if isinstance(content_object, AntrSitzung):
            sitzung = content_object
            daten = sitzung.gesammelte_daten or {}
            # Felddefinitionen aus allen Schritten sammeln für Labels
            label_map = {}
            for schritt in sitzung.pfad.schritte.all():
                for feld in (schritt.felder_json or []):
                    fid = feld.get("id") or feld.get("feldId")
                    if fid:
                        label_map[fid] = feld.get("label") or fid
            for schluessel, wert in daten.items():
                label = label_map.get(schluessel, schluessel)
                antrag_felder.append({"label": label, "wert": wert, "schluessel": schluessel})
    except Exception:
        pass

    # Briefvorlagen gefiltert nach Pfad-Kuerzel
    vorlagen = []
    briefe = []
    dateien = []
    try:
        from korrespondenz.models import Briefvorlage, Briefvorgang
        kuerzel = sitzung.pfad.kuerzel.upper() if sitzung else ""
        if kuerzel:
            vorlagen = list(Briefvorlage.objects.filter(ist_aktiv=True, pfad_kuerzel__iexact=kuerzel).order_by("gruppe", "titel"))
        if not vorlagen:
            vorlagen = list(Briefvorlage.objects.filter(ist_aktiv=True).order_by("titel"))
        briefe = list(Briefvorgang.objects.filter(workflow_task=task).select_related("vorlage", "bearbeiter_signiert_von").order_by("-erstellt_am"))
    except Exception:
        pass
    try:
        if sitzung:
            dateien = list(sitzung.dateien.all())
    except Exception:
        pass

    # Verteiler-Empfaenger laden
    verteiler = []
    try:
        from post.models import VerteilEmpfaenger
        verteiler = list(VerteilEmpfaenger.objects.filter(workflow_task=task))
    except Exception:
        pass

    kontext = {
        "task": task,
        "content_object": content_object,
        "sitzung": sitzung,
        "antrag_felder": antrag_felder,
        "ausgaenge": list(ausgaenge),
        "hat_ausgaenge": hat_ausgaenge,
        "entscheidung_choices": WorkflowTask.ENTSCHEIDUNG_CHOICES,
        "vorlagen": vorlagen,
        "briefe": briefe,
        "dateien": dateien,
        "verteiler": verteiler,
    }
    return render(request, "workflow/task_detail.html", kontext)


@login_required
@require_POST
def task_abholen(request, pk):
    """Claim einen Gruppen-Task (nur fuer Gruppen-Tasks)."""
    task = get_object_or_404(WorkflowTask, pk=pk)
    user = request.user

    if task.zugewiesen_an_gruppe and user.groups.filter(pk=task.zugewiesen_an_gruppe.pk).exists():
        if task.claimed_von is None:
            from django.utils import timezone
            task.claimed_von = user
            task.claimed_am = timezone.now()
            task.status = WorkflowTask.STATUS_IN_BEARBEITUNG
            task.save(update_fields=["claimed_von", "claimed_am", "status"])
            from core.models import audit
            audit(
                request,
                aktion="geaendert",
                app="workflow",
                objekt_typ="WorkflowTask",
                objekt_id=task.pk,
                beschreibung=f"Task '{task.step.titel}' abgeholt aus Gruppen-Queue ({task.zugewiesen_an_gruppe.name})",
            )
            messages.success(request, f"Task '{task.step.titel}' wurde abgeholt.")
        else:
            messages.warning(request, "Dieser Task wurde bereits von jemand anderem abgeholt.")
    else:
        messages.error(request, "Du bist nicht Mitglied der zustaendigen Gruppe.")

    return redirect("workflow:task_detail", pk=task.pk)


@login_required
@require_POST
def task_zurueckgeben(request, pk):
    """Geclaimed-Task zurueck in den Gruppen-Pool geben."""
    task = get_object_or_404(WorkflowTask, pk=pk, claimed_von=request.user)
    task.claimed_von = None
    task.claimed_am = None
    task.status = WorkflowTask.STATUS_OFFEN
    task.save(update_fields=["claimed_von", "claimed_am", "status"])
    messages.info(request, "Task wurde zurueckgegeben.")
    return redirect("workflow:arbeitsstapel")


# ---------------------------------------------------------------------------
# Workflow-Template Liste & Detail
# ---------------------------------------------------------------------------


@login_required
def workflow_liste(request):
    """Liste aller aktiven Workflow-Templates."""
    templates = WorkflowTemplate.objects.filter(ist_aktiv=True).annotate(
        anzahl_instanzen=Count("instanzen"),
        laufende_instanzen=Count(
            "instanzen", filter=Q(instanzen__status=WorkflowInstance.STATUS_LAUFEND)
        ),
    )
    return render(request, "workflow/liste.html", {"templates": templates})


@login_required
def workflow_detail(request, pk):
    """Detail-Ansicht eines Workflow-Templates."""
    template = get_object_or_404(WorkflowTemplate, pk=pk)
    schritte = template.schritte.select_related(
        "zustaendig_gruppe", "zustaendig_user"
    ).order_by("reihenfolge")
    transitions = template.transitions.select_related(
        "von_schritt", "zu_schritt"
    ).order_by("prioritaet")
    instanzen = template.instanzen.select_related(
        "gestartet_von", "aktueller_schritt"
    ).order_by("-gestartet_am")[:20]

    kontext = {
        "template": template,
        "schritte": schritte,
        "transitions": transitions,
        "instanzen": instanzen,
    }
    return render(request, "workflow/detail.html", kontext)


# ---------------------------------------------------------------------------
# Workflow-Instanz Detail
# ---------------------------------------------------------------------------


@login_required
def instanz_detail(request, pk):
    """Detail-Ansicht einer laufenden Workflow-Instanz."""
    instanz = get_object_or_404(
        WorkflowInstance.objects.select_related(
            "template", "gestartet_von", "aktueller_schritt"
        ),
        pk=pk,
    )
    tasks = instanz.tasks.select_related(
        "step", "zugewiesen_an_gruppe", "zugewiesen_an_user", "erledigt_von"
    ).order_by("step__reihenfolge", "erstellt_am")

    kontext = {
        "instanz": instanz,
        "tasks": tasks,
        "content_object": instanz.content_object,
    }
    return render(request, "workflow/instanz_detail.html", kontext)


# ---------------------------------------------------------------------------
# Visueller Workflow-Editor
# ---------------------------------------------------------------------------


@login_required
def workflow_editor(request, pk=None):
    """Visueller Workflow-Editor mit vis.js.

    Beim Aufruf ohne pk wird ein neues leeres Template angelegt.
    """
    if not request.user.is_staff:
        messages.error(request, "Nur Administratoren duerfen Workflows bearbeiten.")
        return redirect("workflow:liste")

    if pk:
        template = get_object_or_404(WorkflowTemplate, pk=pk)
    else:
        template = None

    # Alle Gruppen fuer Zustaendigkeit
    from django.contrib.auth.models import Group
    gruppen = list(Group.objects.values("id", "name").order_by("name"))

    # Alle User fuer spezifische Zustaendigkeit
    from django.contrib.auth import get_user_model
    User = get_user_model()
    users = list(User.objects.filter(is_active=True).values("id", "username", "first_name", "last_name"))

    kontext = {
        "template": template,
        "gruppen_json": gruppen,
        "users_json": users,
        "aktion_choices": WorkflowStep.AKTION_CHOICES,
        "rolle_choices": WorkflowStep.ROLLE_CHOICES,
        "schritt_typ_choices": WorkflowStep.SCHRITT_TYP_CHOICES,
        "bedingung_choices": WorkflowTransition.BEDINGUNG_CHOICES,
        "entscheidung_choices": WorkflowTransition.ENTSCHEIDUNG_CHOICES,
    }
    return render(request, "workflow/workflow_editor.html", kontext)


@login_required
@require_POST
def workflow_editor_save(request):
    """Speichert den Workflow-Graphen (JSON) aus dem Editor."""
    if not request.user.is_staff:
        return JsonResponse({"error": "Keine Berechtigung"}, status=403)

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError) as exc:
        return JsonResponse({"error": f"Ungueltige JSON-Daten: {exc}"}, status=400)

    template_id = data.get("template_id")
    meta = data.get("meta", {})
    nodes = data.get("nodes", [])
    edges = data.get("edges", [])

    if template_id:
        template = get_object_or_404(WorkflowTemplate, pk=template_id)
    else:
        template = WorkflowTemplate(erstellt_von=request.user)

    template.name = meta.get("name", template.name or "Neuer Workflow")
    template.beschreibung = meta.get("beschreibung", template.beschreibung or "")
    template.kategorie = meta.get("kategorie", template.kategorie or WorkflowTemplate.KATEGORIE_GENEHMIGUNG)
    template.ist_graph_workflow = True
    template.ist_aktiv = meta.get("ist_aktiv", True)
    template.save()

    # Bestehende Schritte und Transitions loeschen
    template.schritte.all().delete()
    # Transitions werden per CASCADE geloescht

    # Neue Schritte anlegen
    node_map = {}  # node_id (Editor) -> WorkflowStep
    from django.contrib.auth.models import Group
    from django.contrib.auth import get_user_model
    User = get_user_model()

    for i, node in enumerate(nodes):
        node_id = node.get("id", "")
        gruppe_id = node.get("gruppeId")
        user_id = node.get("userId")
        eskalation_id = node.get("eskalationGruppeId")

        gruppe = None
        if gruppe_id:
            try:
                gruppe = Group.objects.get(pk=gruppe_id)
            except Group.DoesNotExist:
                pass

        zustaendig_user = None
        if user_id:
            try:
                zustaendig_user = User.objects.get(pk=user_id)
            except User.DoesNotExist:
                pass

        eskalation_gruppe = None
        if eskalation_id:
            try:
                eskalation_gruppe = Group.objects.get(pk=eskalation_id)
            except Group.DoesNotExist:
                pass

        schritt = WorkflowStep.objects.create(
            template=template,
            node_id=str(node_id),
            titel=node.get("label", "Schritt"),
            beschreibung=node.get("beschreibung", ""),
            schritt_typ=node.get("schritt_typ", "task"),
            aktion_typ=node.get("aktion_typ", WorkflowStep.AKTION_PRUEFEN),
            zustaendig_rolle=node.get("zustaendig_rolle", WorkflowStep.ROLLE_GRUPPE),
            zustaendig_gruppe=gruppe,
            zustaendig_user=zustaendig_user,
            eskalation_an_gruppe=eskalation_gruppe,
            eskalation_nach_tagen=node.get("eskalation_nach_tagen", 0),
            frist_tage=node.get("frist_tage", 3),
            reihenfolge=i + 1,
            pos_x=node.get("x", 200),
            pos_y=node.get("y", 200),
            auto_config=node.get("auto_config"),
        )
        node_map[str(node_id)] = schritt

    # Transitions anlegen
    for edge in edges:
        from_id = str(edge.get("from", ""))
        to_id = str(edge.get("to", "")) if edge.get("to") else None
        von = node_map.get(from_id)
        zu = node_map.get(to_id) if to_id else None
        if not von:
            continue

        WorkflowTransition.objects.create(
            template=template,
            von_schritt=von,
            zu_schritt=zu,
            bedingung_typ=edge.get("bedingung_typ", "immer"),
            bedingung_entscheidung=edge.get("bedingung_entscheidung") or None,
            bedingung_feld=edge.get("bedingung_feld", ""),
            bedingung_operator=edge.get("bedingung_operator", ""),
            bedingung_wert=edge.get("bedingung_wert", ""),
            label=edge.get("label", ""),
            prioritaet=edge.get("prioritaet", 1),
        )

    return JsonResponse({"status": "ok", "template_id": template.pk})


@login_required
def workflow_editor_load(request, pk):
    """Laedt die Graph-Daten eines Templates als JSON fuer den Editor."""
    template = get_object_or_404(WorkflowTemplate, pk=pk)
    schritte = template.schritte.select_related(
        "zustaendig_gruppe", "zustaendig_user"
    ).order_by("reihenfolge")
    transitions = template.transitions.select_related(
        "von_schritt", "zu_schritt"
    ).order_by("prioritaet")

    nodes = []
    for s in schritte:
        nodes.append({
            "id": s.node_id or str(s.pk),
            "label": s.titel,
            "beschreibung": s.beschreibung,
            "schritt_typ": s.schritt_typ,
            "aktion_typ": s.aktion_typ,
            "zustaendig_rolle": s.zustaendig_rolle,
            "gruppeId": s.zustaendig_gruppe_id,
            "gruppeName": s.zustaendig_gruppe.name if s.zustaendig_gruppe else "",
            "userId": s.zustaendig_user_id,
            "userName": (
                s.zustaendig_user.get_full_name() or s.zustaendig_user.username
                if s.zustaendig_user else ""
            ),
            "eskalationGruppeId": s.eskalation_an_gruppe_id,
            "eskalation_nach_tagen": s.eskalation_nach_tagen,
            "frist_tage": s.frist_tage,
            "auto_config": s.auto_config,
            "x": s.pos_x,
            "y": s.pos_y,
        })

    edges = []
    for t in transitions:
        edges.append({
            "from": t.von_schritt.node_id or str(t.von_schritt.pk),
            "to": (t.zu_schritt.node_id or str(t.zu_schritt.pk)) if t.zu_schritt else None,
            "bedingung_typ": t.bedingung_typ,
            "bedingung_entscheidung": t.bedingung_entscheidung or "",
            "bedingung_feld": t.bedingung_feld,
            "bedingung_operator": t.bedingung_operator,
            "bedingung_wert": t.bedingung_wert,
            "label": t.label,
            "prioritaet": t.prioritaet,
        })

    return JsonResponse({
        "template_id": template.pk,
        "meta": {
            "name": template.name,
            "beschreibung": template.beschreibung,
            "kategorie": template.kategorie,
            "ist_aktiv": template.ist_aktiv,
        },
        "nodes": nodes,
        "edges": edges,
    })


# ---------------------------------------------------------------------------
# Prozessantrag
# ---------------------------------------------------------------------------


@login_required
def prozessantrag_erstellen(request):
    """Antrag auf Erstellung eines neuen Prozesses einreichen."""
    form = ProzessAntragForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        antrag = form.save(commit=False)
        antrag.antragsteller = request.user
        antrag.save()
        messages.success(
            request,
            "Dein Prozessantrag wurde eingereicht. Wir melden uns bei dir.",
        )
        return redirect("workflow:prozessantrag_liste")

    return render(request, "workflow/prozessantrag_erstellen.html", {"form": form})


@login_required
def prozessantrag_liste(request):
    """Liste der eigenen Prozessantraege."""
    antraege = ProzessAntrag.objects.filter(antragsteller=request.user).order_by("-erstellt_am")
    return render(request, "workflow/prozessantrag_liste.html", {"antraege": antraege})


@login_required
def prozesszentrale(request):
    """Admin-Uebersicht fuer Prozessverantwortliche.

    Zeigt alle Prozessantraege, laufende Instanzen und Trigger.
    """
    if not request.user.is_staff:
        messages.error(request, "Keine Berechtigung.")
        return redirect("workflow:arbeitsstapel")

    antraege = ProzessAntrag.objects.select_related("antragsteller").order_by("-erstellt_am")
    instanzen_laufend = WorkflowInstance.objects.filter(
        status=WorkflowInstance.STATUS_LAUFEND
    ).select_related("template", "gestartet_von", "aktueller_schritt").order_by("-gestartet_am")[:50]
    ueberfaellige_tasks = WorkflowTask.objects.filter(
        status__in=[WorkflowTask.STATUS_OFFEN, WorkflowTask.STATUS_IN_BEARBEITUNG],
    ).select_related(
        "step__template", "zugewiesen_an_gruppe", "zugewiesen_an_user"
    ).order_by("frist")

    kontext = {
        "antraege": antraege,
        "instanzen_laufend": instanzen_laufend,
        "ueberfaellige_tasks": [t for t in ueberfaellige_tasks if t.ist_ueberfaellig],
        "anzahl_offen": ueberfaellige_tasks.count(),
    }
    return render(request, "workflow/prozesszentrale.html", kontext)


# ---------------------------------------------------------------------------
# Trigger-Verwaltung
# ---------------------------------------------------------------------------


@login_required
def trigger_liste(request):
    """Liste aller konfigurierten Workflow-Trigger."""
    if not request.user.is_staff:
        messages.error(request, "Keine Berechtigung.")
        return redirect("workflow:arbeitsstapel")

    trigger = WorkflowTrigger.objects.select_related("content_type").order_by("name")
    return render(request, "workflow/trigger_liste.html", {"trigger": trigger})


@login_required
def trigger_erstellen(request):
    """Neuen Workflow-Trigger anlegen."""
    if not request.user.is_staff:
        messages.error(request, "Keine Berechtigung.")
        return redirect("workflow:arbeitsstapel")

    form = WorkflowTriggerForm(request.POST or None)
    if request.method == "POST" and form.is_valid():
        form.save()
        messages.success(request, "Trigger wurde angelegt.")
        return redirect("workflow:trigger_liste")

    return render(request, "workflow/trigger_form.html", {"form": form, "aktion": "Erstellen"})


@login_required
def trigger_bearbeiten(request, pk):
    """Bestehenden Trigger bearbeiten."""
    if not request.user.is_staff:
        messages.error(request, "Keine Berechtigung.")
        return redirect("workflow:arbeitsstapel")

    trigger = get_object_or_404(WorkflowTrigger, pk=pk)
    form = WorkflowTriggerForm(request.POST or None, instance=trigger)
    if request.method == "POST" and form.is_valid():
        form.save()
        messages.success(request, "Trigger wurde gespeichert.")
        return redirect("workflow:trigger_liste")

    return render(request, "workflow/trigger_form.html", {"form": form, "aktion": "Bearbeiten", "trigger": trigger})


@login_required
@require_POST
def trigger_loeschen(request, pk):
    """Trigger loeschen."""
    if not request.user.is_staff:
        messages.error(request, "Keine Berechtigung.")
        return redirect("workflow:arbeitsstapel")

    trigger = get_object_or_404(WorkflowTrigger, pk=pk)
    name = trigger.name
    trigger.delete()
    messages.success(request, f"Trigger '{name}' wurde geloescht.")
    return redirect("workflow:trigger_liste")


# ---------------------------------------------------------------------------
# API: ContentType-Liste fuer Trigger-Formular
# ---------------------------------------------------------------------------


@login_required
def api_content_types(request):
    """Gibt alle ContentTypes als JSON zurueck (fuer Trigger-Formular)."""
    if not request.user.is_staff:
        return JsonResponse({"error": "Keine Berechtigung"}, status=403)

    cts = ContentType.objects.all().values("id", "app_label", "model").order_by("app_label", "model")
    return JsonResponse({"content_types": list(cts)})


# ---------------------------------------------------------------------------
# Workflow-Status fuer ein Content-Objekt (einbettbar in andere Apps)
# ---------------------------------------------------------------------------


@login_required
def workflow_status_partial(request, content_type_id, object_id):
    """Gibt den Workflow-Status eines Objekts als Partial zurueck (fuer HTMX).

    Kann in beliebige Templates eingebettet werden.
    """
    content_type = get_object_or_404(ContentType, pk=content_type_id)
    instanzen = WorkflowInstance.objects.filter(
        content_type=content_type,
        object_id=object_id,
    ).select_related("template", "aktueller_schritt").order_by("-gestartet_am")

    offene_tasks = WorkflowTask.objects.filter(
        instance__in=instanzen,
        status__in=[WorkflowTask.STATUS_OFFEN, WorkflowTask.STATUS_IN_BEARBEITUNG],
    ).select_related("step", "zugewiesen_an_gruppe")

    kontext = {
        "instanzen": instanzen,
        "offene_tasks": offene_tasks,
    }
    return render(request, "workflow/partials/_workflow_status.html", kontext)


# ---------------------------------------------------------------------------
# Task-Dokumente: Hilfsfunktionen
# ---------------------------------------------------------------------------

def _docx_zu_pdf(docx_bytes: bytes, brief) -> bytes:
    """Konvertiert DOCX-Bytes zu PDF via python-docx → HTML → WeasyPrint."""
    import io
    from docx import Document
    from weasyprint import HTML

    doc = Document(io.BytesIO(docx_bytes))

    def _escape(t):
        return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    zeilen = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1") or style.startswith("berschrift 1"):
            zeilen.append(f"<h1>{_escape(text)}</h1>")
        elif style.startswith("Heading 2") or style.startswith("berschrift 2"):
            zeilen.append(f"<h2>{_escape(text)}</h2>")
        elif text:
            zeilen.append(f"<p>{_escape(text)}</p>")
        else:
            zeilen.append("<p>&nbsp;</p>")

    for table in doc.tables:
        zeilen.append('<table border="1" cellpadding="4" cellspacing="0" style="border-collapse:collapse;width:100%;margin:8px 0;">')
        for row in table.rows:
            zeilen.append("<tr>")
            for cell in row.cells:
                cell_text = " ".join(p.text for p in cell.paragraphs if p.text.strip())
                zeilen.append(f"<td>{_escape(cell_text)}</td>")
            zeilen.append("</tr>")
        zeilen.append("</table>")

    absender = getattr(brief, "absender_name", "") or ""
    betreff  = getattr(brief, "betreff", "") or ""
    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; font-size: 11pt; margin: 2cm; }}
  h1 {{ font-size: 14pt; }} h2 {{ font-size: 12pt; }}
  p {{ margin: 4px 0; line-height: 1.5; }}
</style></head><body>
{''.join(zeilen)}
</body></html>"""

    return HTML(string=html).write_pdf()


# ---------------------------------------------------------------------------
# Task-Dokumente: Schreiben erstellen + signieren
# ---------------------------------------------------------------------------

@login_required
@require_POST
def workflow_brief_erstellen(request, task_pk):
    """Erstellt einen Briefvorgang aus einer Vorlage, befuellt mit Antragsdaten."""
    task = get_object_or_404(WorkflowTask, pk=task_pk)
    if not task.kann_bearbeiten(request.user) and not request.user.is_staff:
        messages.error(request, "Kein Zugriff.")
        return redirect("workflow:task_detail", pk=task_pk)

    vorlage_pk = request.POST.get("vorlage_pk")
    from korrespondenz.models import Briefvorlage, Briefvorgang
    vorlage = get_object_or_404(Briefvorlage, pk=vorlage_pk, ist_aktiv=True)

    # Antragsdaten aus verknuepfter Sitzung
    sitzung = None
    content_object = task.instance.content_object
    try:
        from formulare.models import AntrSitzung
        if isinstance(content_object, AntrSitzung):
            sitzung = content_object
    except Exception:
        pass

    # Empfaenger und Platzhalter aus Antragsdaten ableiten
    empfaenger_name = ""
    platz_extra = {}
    if sitzung:
        daten = sitzung.gesammelte_daten or {}
        # Feldlabels fuer Mapping aufbauen
        for schritt in sitzung.pfad.schritte.all():
            for feld in (schritt.felder_json or []):
                fid = feld.get("id") or feld.get("feldId", "")
                label = feld.get("label", "")
                wert = daten.get(fid, "")
                if fid:
                    platz_extra[fid] = str(wert) if wert else ""
                if label:
                    # Auch normalisiertes Label als Platzhalter
                    label_key = label.lower().replace(" ", "_").replace("-", "_")
                    platz_extra[label_key] = str(wert) if wert else ""
                # Empfaenger-Name aus typischen Feldern
                if not empfaenger_name and label.lower() in ("name", "vor- und nachname", "vor und nachname", "vollstaendiger name"):
                    empfaenger_name = str(wert)

        platz_extra["vorgangsnummer"] = sitzung.vorgangsnummer or ""
        platz_extra["antrag_datum"] = sitzung.abgeschlossen_am.strftime("%d.%m.%Y") if sitzung.abgeschlossen_am else ""
        platz_extra["pfad_name"] = sitzung.pfad.name
        if sitzung.user:
            platz_extra["antragsteller"] = sitzung.user.get_full_name() or sitzung.user.username
            if not empfaenger_name:
                empfaenger_name = sitzung.user.get_full_name() or sitzung.user.username

    from korrespondenz.models import Firmendaten
    firma = Firmendaten.laden()
    from django.utils import timezone as tz

    vorgang = Briefvorgang.objects.create(
        vorlage=vorlage,
        sitzung=sitzung,
        workflow_task=task,
        absender_name=vorlage.default_absender_name or firma.firmenname,
        absender_strasse=vorlage.default_absender_strasse or firma.strasse,
        absender_ort=vorlage.default_absender_ort or firma.plz_ort,
        absender_telefon=vorlage.default_absender_telefon or firma.telefon,
        absender_email=vorlage.default_absender_email or firma.email,
        empfaenger_name=empfaenger_name or "—",
        ort=vorlage.default_ort or firma.ort or "",
        datum=tz.localdate(),
        betreff=vorlage.titel,
        grussformel=vorlage.default_grussformel or firma.grussformel,
        unterschrift_name=request.user.get_full_name() or request.user.username,
        erstellt_von=request.user,
    )

    # Vorlage befuellen
    try:
        from korrespondenz.views import _erstelle_platzhalter, _fuelle_vorlage
        platz = _erstelle_platzhalter(vorgang)
        platz.update(platz_extra)
        ausgefuellt = _fuelle_vorlage(bytes(vorlage.inhalt), platz)
        Briefvorgang.objects.filter(pk=vorgang.pk).update(inhalt=ausgefuellt)
        vorgang.inhalt = ausgefuellt
    except Exception as exc:
        logger.error("Vorlage befuellen fehlgeschlagen: %s", exc)
        messages.warning(request, "Vorlage konnte nicht vollstaendig befuellt werden.")

    messages.success(request, f"Schreiben '{vorlage.titel}' erstellt.")
    return redirect(f"/korrespondenz/{vorgang.pk}/editor/?zurueck=/workflow/task/{task_pk}/")


@login_required
@require_POST
def workflow_brief_signieren(request, brief_pk):
    """Konvertiert Briefvorgang zu PDF und signiert ihn (Bearbeiter-FES)."""
    from korrespondenz.models import Briefvorgang
    brief = get_object_or_404(Briefvorgang, pk=brief_pk)
    task = brief.workflow_task
    if task and not task.kann_bearbeiten(request.user) and not request.user.is_staff:
        return JsonResponse({"ok": False, "fehler": "Kein Zugriff"}, status=403)

    if not brief.inhalt:
        return JsonResponse({"ok": False, "fehler": "Kein Dokument-Inhalt vorhanden"})

    from django.conf import settings as conf
    import json as _json, urllib.request as urlreq
    oo_int = (getattr(conf, "ONLYOFFICE_INTERNAL_URL", "") or getattr(conf, "ONLYOFFICE_URL", "")).rstrip("/")
    base_url = getattr(conf, "WOPI_BASE_URL", getattr(conf, "VORGANGSWERK_BASE_URL", "")).rstrip("/")

    pdf_bytes = None

    # 1. OO ConversionAPI: DOCX → PDF
    if oo_int and base_url:
        try:
            from korrespondenz.views import _oo_jwt
            dl_url = f"{base_url}/korrespondenz/{brief.pk}/oo-download/"
            conv_key = f"conv-brief-{brief.pk}-v{brief.version}"
            payload = {"async": False, "filetype": "docx", "key": conv_key, "outputtype": "pdf", "url": dl_url}
            headers = {"Content-Type": "application/json"}
            secret = getattr(conf, "ONLYOFFICE_JWT_SECRET", "")
            if secret:
                headers["Authorization"] = f"Bearer {_oo_jwt(payload)}"
            req = urlreq.Request(
                f"{oo_int}/ConvertService.ashx",
                data=_json.dumps(payload).encode(),
                headers=headers,
                method="POST",
            )
            with urlreq.urlopen(req, timeout=30) as resp:
                result = _json.loads(resp.read())
            pdf_url = result.get("fileUrl", "")
            if pdf_url:
                if oo_int and getattr(conf, "ONLYOFFICE_URL", ""):
                    pdf_url = pdf_url.replace(getattr(conf, "ONLYOFFICE_URL", "").rstrip("/"), oo_int, 1)
                dl_req = urlreq.Request(pdf_url)
                if secret:
                    import jwt as pyjwt
                    dl_req.add_header("Authorization", f"Bearer {pyjwt.encode({'url': pdf_url}, secret, algorithm='HS256')}")
                with urlreq.urlopen(dl_req, timeout=30) as r:
                    pdf_bytes = r.read()
        except Exception as exc:
            logger.error("OO-Konvertierung fehlgeschlagen fuer Brief %s: %s", brief_pk, exc)

    # Fallback: DOCX → HTML → WeasyPrint wenn OO nicht erreichbar
    if not pdf_bytes:
        try:
            pdf_bytes = _docx_zu_pdf(bytes(brief.inhalt), brief)
        except Exception as exc:
            logger.error("Fallback-PDF-Konvertierung fehlgeschlagen fuer Brief %s: %s", brief_pk, exc)
            return JsonResponse({"ok": False, "fehler": f"PDF-Konvertierung fehlgeschlagen: {exc}"})

    # 2. FES-Signatur
    try:
        from signatur.services import signiere_pdf
        signiert = signiere_pdf(pdf_bytes, request.user, brief.betreff or "Schreiben")
    except Exception as exc:
        logger.error("FES-Signatur fehlgeschlagen fuer Brief %s: %s", brief_pk, exc)
        return JsonResponse({"ok": False, "fehler": f"Signatur fehlgeschlagen: {exc}"})

    from django.utils import timezone as tz
    Briefvorgang.objects.filter(pk=brief_pk).update(
        signiert_pdf=signiert,
        bearbeiter_signiert_von=request.user,
        bearbeiter_signiert_am=tz.now(),
        status=Briefvorgang.STATUS_FERTIG,
    )
    return JsonResponse({"ok": True})
