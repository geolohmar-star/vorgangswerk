"""Korrespondenz-App Views: Briefvorlagen, Briefvorgaenge, Collabora WOPI."""
import io
import logging

from django.conf import settings as django_settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.csrf import csrf_exempt

from .forms import BriefvorgangForm
from .models import Briefvorlage, Briefvorgang

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Hilfsfunktionen: DOCX-Befuellung
# ---------------------------------------------------------------------------

def _fuelle_vorlage(vorlage_bytes: bytes, platzhalter: dict) -> bytes:
    """Ersetzt alle {{schluessel}}-Platzhalter in der DOCX-Vorlage.

    Verarbeitet normale Absaetze, Tabellenzellen sowie Kopf- und Fusszeilen.
    Zusammengesetzte Platzhalter (ueber mehrere Runs gesplittet) werden korrekt
    behandelt, indem der gesamte Absatz-Text rekonstruiert wird.
    """
    from docx import Document

    doc = Document(io.BytesIO(vorlage_bytes))

    def _ersetze_absatz(para):
        full_text = "".join(run.text for run in para.runs)
        neuer_text = full_text
        for schluessel, wert in platzhalter.items():
            token = "{{" + schluessel + "}}"
            neuer_text = neuer_text.replace(token, str(wert) if wert else "")
        if neuer_text != full_text and para.runs:
            para.runs[0].text = neuer_text
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        _ersetze_absatz(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _ersetze_absatz(para)

    for section in doc.sections:
        for para in section.footer.paragraphs:
            _ersetze_absatz(para)
        for para in section.header.paragraphs:
            _ersetze_absatz(para)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _datum_deutsch(datum) -> str:
    monate = ["", "Januar", "Februar", "Maerz", "April", "Mai", "Juni",
              "Juli", "August", "September", "Oktober", "November", "Dezember"]
    return f"{datum.day}. {monate[datum.month]} {datum.year}"


def _erstelle_platzhalter(vorgang: Briefvorgang) -> dict:
    """Erstellt das Platzhalter-Dict aus einem Briefvorgang."""
    from .models import Firmendaten
    org_email = getattr(django_settings, "EMAIL_HOST_USER", "")
    firma = Firmendaten.laden()
    ort = vorgang.ort or firma.ort or ""

    return {
        "absender_name":    vorgang.absender_name    or firma.firmenname,
        "absender_strasse": vorgang.absender_strasse or firma.strasse,
        "absender_ort":     vorgang.absender_ort     or firma.plz_ort,
        "absender_telefon": vorgang.absender_telefon or firma.telefon,
        "absender_email":   vorgang.absender_email   or firma.email or org_email,
        "empfaenger_name":    vorgang.empfaenger_name,
        "empfaenger_zusatz":  vorgang.empfaenger_zusatz,
        "empfaenger_strasse": vorgang.empfaenger_strasse,
        "empfaenger_plz_ort": vorgang.empfaenger_plz_ort,
        "empfaenger_land":    vorgang.empfaenger_land,
        "ort_datum":          f"{ort}, {_datum_deutsch(vorgang.datum)}",
        "datum":              _datum_deutsch(vorgang.datum),
        "betreff":            vorgang.betreff,
        "anrede":             vorgang.anrede,
        "brieftext":          vorgang.brieftext,
        "grussformel":        vorgang.grussformel,
        "unterschrift_name":  vorgang.unterschrift_name,
        "unterschrift_titel": vorgang.unterschrift_titel,
        "fusszeile_firmenname": vorgang.vorlage.fusszeile_firmenname,
        "fusszeile_telefon":    vorgang.vorlage.fusszeile_telefon,
        "fusszeile_telefax":    vorgang.vorlage.fusszeile_telefax,
        "fusszeile_email":      vorgang.vorlage.fusszeile_email or org_email,
        "fusszeile_internet":   vorgang.vorlage.fusszeile_internet,
    }


def _get_oo_url():
    return getattr(django_settings, "ONLYOFFICE_URL", "").rstrip("/")


def _oo_jwt(payload: dict) -> str:
    """Signiert einen OnlyOffice-Payload als HS256-JWT (via PyJWT)."""
    import jwt
    secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if not secret:
        return ""
    return jwt.encode(payload, secret, algorithm="HS256")


def _oo_jwt_verify(request) -> bool:
    """Prueft den Authorization-Bearer-JWT einer eingehenden OO-Anfrage."""
    import jwt
    secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if not secret:
        return True
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        return False
    try:
        jwt.decode(auth[7:], secret, algorithms=["HS256"])
        return True
    except jwt.PyJWTError:
        return False


_STANDARD_PLATZHALTER = [
    ("absender_name",      "Absender Name"),
    ("absender_strasse",   "Absender Strasse"),
    ("absender_ort",       "Absender PLZ/Ort"),
    ("absender_telefon",   "Absender Telefon"),
    ("absender_email",     "Absender E-Mail"),
    ("empfaenger_name",    "Empfaenger Name"),
    ("empfaenger_zusatz",  "Empfaenger Zusatz"),
    ("empfaenger_strasse", "Empfaenger Strasse"),
    ("empfaenger_plz_ort", "Empfaenger PLZ/Ort"),
    ("empfaenger_land",    "Empfaenger Land"),
    ("ort",                "Ort"),
    ("datum",              "Datum"),
    ("betreff",            "Betreff"),
    ("anrede",             "Anrede"),
    ("brieftext",          "Brieftext"),
    ("grussformel",        "Grussformel"),
    ("unterschrift_name",  "Unterzeichner"),
    ("unterschrift_titel", "Funktion/Titel"),
    ("vorgangsnummer",     "Vorgangsnummer"),
    ("antragsteller",      "Antragsteller"),
    ("antrag_datum",       "Antragsdatum"),
    ("pfad_name",          "Pfad Name"),
    ("fusszeile_firmenname", "Fusszeile Firma"),
    ("fusszeile_telefon",    "Fusszeile Telefon"),
    ("fusszeile_telefax",    "Fusszeile Fax"),
    ("fusszeile_email",      "Fusszeile E-Mail"),
    ("fusszeile_internet",   "Fusszeile Internet"),
]


def _pfad_platzhalter(kuerzel: str) -> list:
    if not kuerzel:
        return []
    from formulare.models import AntrPfad
    pfad = AntrPfad.objects.filter(kuerzel=kuerzel).first()
    if not pfad:
        return []
    result = []
    for schritt in pfad.schritte.all():
        for feld in schritt.felder():
            fid = feld.get("id", "")
            if fid:
                result.append((fid, feld.get("label", fid)))
    return result


# ---------------------------------------------------------------------------
# Views: Briefvorlagen-Verwaltung
# ---------------------------------------------------------------------------

@login_required
def vorlage_liste(request):
    """Listet alle aktiven Briefvorlagen auf, gruppiert nach Pfad und Untergruppe."""
    from django.db.models import Q
    from formulare.models import AntrPfad

    q = request.GET.get("q", "").strip()
    qs = Briefvorlage.objects.filter(ist_aktiv=True).order_by("pfad_kuerzel", "gruppe", "titel")
    if q:
        qs = qs.filter(Q(titel__icontains=q) | Q(pfad_kuerzel__icontains=q) | Q(gruppe__icontains=q))

    pfade = AntrPfad.objects.filter(aktiv=True).order_by("name")
    pfad_namen = {p.kuerzel: p.name for p in pfade if p.kuerzel}

    def _baue_untergruppen(vorlagen):
        untergruppen = {}
        ungrouped = []
        for v in vorlagen:
            if v.gruppe:
                untergruppen.setdefault(v.gruppe, []).append(v)
            else:
                ungrouped.append(v)
        abschnitte = []
        for name in sorted(untergruppen):
            abschnitte.append({"typ": "gruppe", "name": name, "vorlagen": untergruppen[name]})
        for v in ungrouped:
            abschnitte.append({"typ": "einzeln", "vorlage": v})
        return abschnitte

    pfad_gruppen = {}
    allgemein_roh = []
    for v in qs:
        if v.pfad_kuerzel:
            pfad_gruppen.setdefault(v.pfad_kuerzel, []).append(v)
        else:
            allgemein_roh.append(v)

    gruppen_liste = sorted(
        [{"kuerzel": key, "name": pfad_namen.get(key, key),
          "abschnitte": _baue_untergruppen(vorlagen), "anzahl": len(vorlagen)}
         for key, vorlagen in pfad_gruppen.items()],
        key=lambda g: g["name"],
    )

    pfad_vorauswahl = request.GET.get("pfad", "").strip().upper()

    return render(request, "korrespondenz/vorlage_liste.html", {
        "gruppen_liste":    gruppen_liste,
        "allgemein":        _baue_untergruppen(allgemein_roh),
        "allgemein_anzahl": len(allgemein_roh),
        "q":                q,
        "pfade":            pfade,
        "pfad_vorauswahl":  pfad_vorauswahl,
    })


def _generiere_din5008_docx(firma, pfad_felder: list) -> bytes:
    """Erzeugt eine minimale Briefvorlage (nur plain paragraphs) mit Firmendaten
    als Absender und Formular-Feldern als Platzhalter im Brieftext.
    Keine Formatierung (kein Pt/RGBColor/Cm) – Collabora vertraegt nur plain DOCX."""
    import io
    from docx import Document

    doc = Document()

    absender_zeile = ", ".join(filter(None, [
        firma.firmenname, firma.strasse, firma.plz_ort
    ])) or "Absender"
    doc.add_paragraph(absender_zeile)
    doc.add_paragraph("")
    doc.add_paragraph("{{empfaenger_name}}")
    doc.add_paragraph("{{empfaenger_zusatz}}")
    doc.add_paragraph("{{empfaenger_strasse}}")
    doc.add_paragraph("{{empfaenger_plz_ort}}")
    doc.add_paragraph("")

    ort_datum = f"{firma.ort}, {{{{datum}}}}" if firma.ort else "{{ort}}, {{datum}}"
    doc.add_paragraph(ort_datum)
    doc.add_paragraph("")

    betreff_p = doc.add_paragraph()
    betreff_run = betreff_p.add_run("{{betreff}}")
    betreff_run.bold = True

    doc.add_paragraph("{{anrede}}")
    doc.add_paragraph("")

    if pfad_felder:
        for fid, flabel in pfad_felder:
            doc.add_paragraph(f"{flabel}: {{{{{fid}}}}}")
    else:
        doc.add_paragraph("{{brieftext}}")

    doc.add_paragraph("")
    gruss = firma.grussformel or "Mit freundlichen Gruessen"
    doc.add_paragraph(gruss)
    doc.add_paragraph("")
    doc.add_paragraph("{{unterschrift_name}}")
    doc.add_paragraph("{{unterschrift_titel}}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@login_required
def vorlage_generieren(request):
    """Generiert DIN-5008-Vorlage mit Firmendaten + Formular-Platzhaltern."""
    from formulare.models import AntrPfad
    from .models import Firmendaten

    pfad_kuerzel = request.GET.get("pfad", "").strip().upper()
    firma = Firmendaten.laden()

    pfad_felder = []
    pfad_name = ""
    if pfad_kuerzel:
        pfad_obj = AntrPfad.objects.filter(kuerzel=pfad_kuerzel).first()
        if pfad_obj:
            pfad_name = pfad_obj.name
            # Alle Felder des Pfads sammeln
            for schritt in pfad_obj.schritte.all():
                for feld in schritt.felder():
                    fid = feld.get("id", "")
                    if fid:
                        pfad_felder.append((fid, feld.get("label", fid)))

    titel = f"Vorlage – {pfad_name}" if pfad_name else "Geschaeftsbrief"
    inhalt = _generiere_din5008_docx(firma, pfad_felder)

    vorlage = Briefvorlage.objects.create(
        titel               = titel,
        pfad_kuerzel        = pfad_kuerzel,
        inhalt              = inhalt,
        default_absender_name    = firma.firmenname,
        default_absender_strasse = firma.strasse,
        default_absender_ort     = firma.plz_ort,
        default_absender_telefon = firma.telefon,
        default_absender_email   = firma.email,
        default_ort              = firma.ort,
        default_grussformel      = firma.grussformel,
        fusszeile_firmenname     = firma.firmenname,
        fusszeile_telefon        = firma.telefon,
        fusszeile_telefax        = firma.telefax,
        fusszeile_email          = firma.email,
        fusszeile_internet       = firma.internet,
        erstellt_von        = request.user,
    )
    messages.success(request, f'Vorlage "{vorlage.titel}" erstellt – bitte anpassen und speichern.')
    return redirect("korrespondenz:vorlage_editor", pk=vorlage.pk)


@login_required
def firmendaten(request):
    """Globale Firmendaten anzeigen und bearbeiten."""
    from .models import Firmendaten as FD
    firma = FD.laden()
    if request.method == "POST":
        firma.firmenname  = request.POST.get("firmenname", "").strip()
        firma.strasse     = request.POST.get("strasse", "").strip()
        firma.plz_ort     = request.POST.get("plz_ort", "").strip()
        firma.telefon     = request.POST.get("telefon", "").strip()
        firma.telefax     = request.POST.get("telefax", "").strip()
        firma.email       = request.POST.get("email", "").strip()
        firma.internet    = request.POST.get("internet", "").strip()
        firma.ort         = request.POST.get("ort", "").strip()
        firma.grussformel = request.POST.get("grussformel", "").strip()
        firma.save()
        messages.success(request, "Firmendaten gespeichert.")
        return redirect("korrespondenz:firmendaten")
    return render(request, "korrespondenz/firmendaten.html", {"firma": firma})


@login_required
def bankverbindungen(request):
    """Bankverbindungen verwalten (nur Staff)."""
    from .models import Bankverbindung
    if not request.user.is_staff:
        return redirect("korrespondenz:firmendaten")
    banken = Bankverbindung.objects.all()
    return render(request, "korrespondenz/bankverbindungen.html", {"banken": banken})


@login_required
def bankverbindung_speichern(request, pk=0):
    """Bankverbindung anlegen oder bearbeiten (nur Staff)."""
    from .models import Bankverbindung
    if not request.user.is_staff:
        return redirect("korrespondenz:bankverbindungen")
    if pk:
        bank = get_object_or_404(Bankverbindung, pk=pk)
    else:
        bank = Bankverbindung()
    if request.method == "POST":
        bank.kuerzel      = request.POST.get("kuerzel", "").strip()
        bank.bezeichnung  = request.POST.get("bezeichnung", "").strip()
        bank.kontoinhaber = request.POST.get("kontoinhaber", "").strip()
        bank.iban         = request.POST.get("iban", "").strip().replace(" ", "")
        bank.bic          = request.POST.get("bic", "").strip()
        bank.bank_name    = request.POST.get("bank_name", "").strip()
        bank.reihenfolge  = int(request.POST.get("reihenfolge", 0) or 0)
        if bank.kuerzel and bank.bezeichnung:
            bank.save()
            messages.success(request, "Bankverbindung gespeichert.")
        return redirect("korrespondenz:bankverbindungen")
    return render(request, "korrespondenz/bankverbindung_form.html", {"bank": bank})


@login_required
def bankverbindung_loeschen(request, pk):
    """Bankverbindung löschen (nur Staff)."""
    from .models import Bankverbindung
    if not request.user.is_staff:
        return redirect("korrespondenz:bankverbindungen")
    bank = get_object_or_404(Bankverbindung, pk=pk)
    if request.method == "POST":
        bank.delete()
        messages.success(request, "Bankverbindung gelöscht.")
    return redirect("korrespondenz:bankverbindungen")


@login_required
def vorlage_hochladen(request):
    if request.method == "POST":
        titel        = request.POST.get("titel", "").strip()
        beschreibung = request.POST.get("beschreibung", "").strip()
        pfad_kuerzel = request.POST.get("pfad_kuerzel", "").strip().upper()
        gruppe       = request.POST.get("gruppe", "").strip()
        datei        = request.FILES.get("docx_datei")

        if not titel:
            messages.error(request, "Bitte einen Titel angeben.")
            return redirect("korrespondenz:vorlage_liste")
        if not datei or not datei.name.endswith(".docx"):
            messages.error(request, "Bitte eine DOCX-Datei hochladen.")
            return redirect("korrespondenz:vorlage_liste")

        vorlage = Briefvorlage.objects.create(
            titel=titel, beschreibung=beschreibung,
            pfad_kuerzel=pfad_kuerzel, gruppe=gruppe,
            inhalt=datei.read(), erstellt_von=request.user,
        )
        messages.success(request, f'Vorlage "{vorlage.titel}" hochgeladen.')
        return redirect("korrespondenz:vorlage_editor", pk=vorlage.pk)

    return redirect("korrespondenz:vorlage_liste")


@login_required
def vorlage_bearbeiten(request, pk):
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    if request.method == "POST":
        vorlage.titel        = request.POST.get("titel", vorlage.titel).strip()
        vorlage.beschreibung = request.POST.get("beschreibung", "").strip()
        vorlage.pfad_kuerzel = request.POST.get("pfad_kuerzel", "").strip().upper()
        vorlage.gruppe       = request.POST.get("gruppe", "").strip()
        vorlage.save()
        messages.success(request, "Vorlage aktualisiert.")
    return redirect("korrespondenz:vorlage_liste")


@login_required
def vorlage_loeschen(request, pk):
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    if request.method == "POST":
        vorlage.ist_aktiv = False
        vorlage.save()
        messages.success(request, f'Vorlage "{vorlage.titel}" geloescht.')
    return redirect("korrespondenz:vorlage_liste")


@login_required
def vorlage_download(request, pk):
    """Laed die rohe DOCX-Datei einer Vorlage herunter (zur Verifikation)."""
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    dateiname = f"Vorlage_{vorlage.pk}_{vorlage.titel}.docx".replace(" ", "_")
    response = HttpResponse(bytes(vorlage.inhalt), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="{dateiname}"'
    return response


@login_required
def vorlage_klonen(request, pk):
    original = get_object_or_404(Briefvorlage, pk=pk)
    if request.method == "POST":
        titel        = request.POST.get("titel", f"{original.titel} – Kopie").strip()
        pfad_kuerzel = request.POST.get("pfad_kuerzel", original.pfad_kuerzel).strip().upper()
        kopie = Briefvorlage.objects.create(
            titel=titel, beschreibung=original.beschreibung,
            pfad_kuerzel=pfad_kuerzel, gruppe=original.gruppe,
            inhalt=bytes(original.inhalt), erstellt_von=request.user,
        )
        messages.success(request, f'Vorlage "{kopie.titel}" erstellt.')
        return redirect("korrespondenz:vorlage_editor", pk=kopie.pk)
    return redirect("korrespondenz:vorlage_liste")


@login_required
def vorlage_defaults(request, pk):
    """Liefert Standard-Absender-Felder einer Vorlage als JSON (AJAX)."""
    vorlage = get_object_or_404(Briefvorlage, pk=pk, ist_aktiv=True)
    return JsonResponse({
        "absender_name":    vorlage.default_absender_name,
        "absender_strasse": vorlage.default_absender_strasse,
        "absender_ort":     vorlage.default_absender_ort,
        "absender_telefon": vorlage.default_absender_telefon,
        "absender_email":   vorlage.default_absender_email,
        "ort":              vorlage.default_ort,
        "grussformel":      vorlage.default_grussformel,
    })


@login_required
def vorlage_platzhalter(request, pk):
    """Zeigt alle Platzhalter eines Pfads als JSON (fuer Editor-Sidebar)."""
    from formulare.models import AntrPfad
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    felder = []
    if vorlage.pfad_kuerzel:
        pfad = AntrPfad.objects.filter(kuerzel=vorlage.pfad_kuerzel).first()
        if pfad:
            for schritt in pfad.schritte.all():
                for feld in schritt.felder():
                    fid = feld.get("id", "")
                    if fid:
                        felder.append({"id": fid, "label": feld.get("label", fid)})

    standard = [
        {"id": "absender_name",    "label": "Absender Name"},
        {"id": "absender_strasse", "label": "Absender Strasse"},
        {"id": "absender_ort",     "label": "Absender PLZ/Ort"},
        {"id": "absender_telefon", "label": "Absender Telefon"},
        {"id": "absender_email",   "label": "Absender E-Mail"},
        {"id": "empfaenger_name",  "label": "Empfaenger Name"},
        {"id": "empfaenger_strasse","label": "Empfaenger Strasse"},
        {"id": "empfaenger_plz_ort","label": "Empfaenger PLZ/Ort"},
        {"id": "ort_datum",        "label": "Ort, Datum"},
        {"id": "datum",            "label": "Datum"},
        {"id": "betreff",          "label": "Betreff"},
        {"id": "anrede",           "label": "Anrede"},
        {"id": "brieftext",        "label": "Brieftext"},
        {"id": "grussformel",      "label": "Grussformel"},
        {"id": "unterschrift_name",  "label": "Unterzeichner Name"},
        {"id": "unterschrift_titel", "label": "Unterzeichner Titel"},
        {"id": "vorgangsnummer",   "label": "Vorgangsnummer"},
        {"id": "antragsteller",    "label": "Antragsteller"},
        {"id": "antrag_datum",     "label": "Antrag Datum"},
        {"id": "pfad_name",        "label": "Pfad Name"},
    ]
    return JsonResponse({"standard": standard, "pfad_felder": felder})


# ---------------------------------------------------------------------------
# OnlyOffice-Editor fuer Briefvorlagen
# ---------------------------------------------------------------------------

@login_required
def vorlage_editor(request, pk):
    """Oeffnet eine Briefvorlage in OnlyOffice."""
    import time
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    oo_url = _get_oo_url()
    if not oo_url:
        messages.error(request, "OnlyOffice nicht konfiguriert (ONLYOFFICE_URL fehlt).")
        return redirect("korrespondenz:vorlage_liste")

    base_url = getattr(django_settings, "WOPI_BASE_URL",
                       getattr(django_settings, "VORGANGSWERK_BASE_URL",
                               request.build_absolute_uri("/").rstrip("/")))

    doc_key = f"vorlage-{vorlage.pk}-t{int(time.time())}"
    Briefvorlage.objects.filter(pk=pk).update(doc_key=doc_key)

    config = {
        "document": {
            "fileType":     "docx",
            "key":          doc_key,
            "title":        f"Vorlage_{vorlage.titel}.docx",
            "url":          f"{base_url}/korrespondenz/vorlagen/{pk}/oo-download/",
        },
        "documentType": "word",
        "editorConfig": {
            "callbackUrl": f"{base_url}/korrespondenz/vorlagen/{pk}/oo-callback/",
            "lang":        "de-DE",
            "region":      "de",
            "mode":        "edit",
            "user": {
                "id":   str(request.user.pk),
                "name": request.user.get_full_name() or request.user.username,
            },
            "customization": {"spellcheck": True},
        },
    }
    token = _oo_jwt(config)

    response = render(request, "korrespondenz/oo_editor.html", {
        "titel":               vorlage.titel,
        "zurueck_url":         "/korrespondenz/vorlagen/",
        "oo_url":              oo_url,
        "oo_config":           config,
        "oo_jwt":              token,
        "vorlage":             vorlage,
        "ist_vorlage":         True,
        "standard_platzhalter": _STANDARD_PLATZHALTER,
        "pfad_platzhalter":    _pfad_platzhalter(vorlage.pfad_kuerzel),
    })
    response["Cache-Control"] = "no-store"
    return response


def vorlage_oo_download(request, pk):
    """Liefert die DOCX-Datei an OnlyOffice (JWT-gesichert via Authorization-Header)."""
    if not _oo_jwt_verify(request):
        return HttpResponse("Unauthorized", status=401)
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    return HttpResponse(
        bytes(vorlage.inhalt),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@csrf_exempt
def vorlage_oo_callback(request, pk):
    """OnlyOffice Speicher-Callback fuer Briefvorlagen."""
    import json, urllib.request as urlreq, jwt as pyjwt
    if request.method != "POST":
        return JsonResponse({"error": 0})
    if not _oo_jwt_verify(request):
        return JsonResponse({"error": 1})
    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": 1})

    status = data.get("status")
    if status not in (2, 6):
        return JsonResponse({"error": 0})

    download_url = data.get("url", "")
    if not download_url:
        return JsonResponse({"error": 0})

    # Oeffentliche OO-URL auf interne umschreiben
    oo_pub = getattr(django_settings, "ONLYOFFICE_URL", "").rstrip("/")
    oo_int = getattr(django_settings, "ONLYOFFICE_INTERNAL_URL", "").rstrip("/")
    if oo_pub and oo_int and download_url.startswith(oo_pub):
        download_url = download_url.replace(oo_pub, oo_int, 1)

    try:
        req = urlreq.Request(download_url)
        secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
        if secret:
            dl_token = pyjwt.encode({"url": download_url}, secret, algorithm="HS256")
            req.add_header("Authorization", f"Bearer {dl_token}")
        with urlreq.urlopen(req, timeout=30) as resp:
            inhalt = resp.read()
    except Exception as exc:
        logger.error("OO-Callback Vorlage %s Download-Fehler: %s", pk, exc)
        return JsonResponse({"error": 1})

    from django.db.models import F
    Briefvorlage.objects.filter(pk=pk).update(inhalt=inhalt, version=F("version") + 1)
    return JsonResponse({"error": 0})


@login_required
def vorlage_oo_forcesave(request, pk):
    """Loest Forcesave im OnlyOffice CommandService aus."""
    import json, urllib.request as urlreq
    if request.method != "POST":
        return JsonResponse({"ok": False})
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    doc_key = vorlage.doc_key
    if not doc_key:
        return JsonResponse({"ok": False, "fehler": "Kein Doc-Key – Seite neu laden."})

    oo_int = (getattr(django_settings, "ONLYOFFICE_INTERNAL_URL", "")
              or getattr(django_settings, "ONLYOFFICE_URL", "")).rstrip("/")
    if not oo_int:
        return JsonResponse({"ok": False, "fehler": "OnlyOffice nicht konfiguriert"})

    payload = {"c": "forcesave", "key": doc_key}
    headers = {"Content-Type": "application/json"}
    secret  = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
    if secret:
        headers["Authorization"] = f"Bearer {_oo_jwt(payload)}"

    try:
        req = urlreq.Request(
            f"{oo_int}/coauthoring/CommandService.ashx",
            data=json.dumps(payload).encode(),
            headers=headers,
            method="POST",
        )
        with urlreq.urlopen(req, timeout=10) as resp:
            antwort = json.loads(resp.read())
        if antwort.get("error", 1) not in (0, 4):
            return JsonResponse({"ok": False, "fehler": str(antwort)})
    except Exception as exc:
        logger.error("Vorlage %s Forcesave-Fehler: %s", pk, exc)
        return JsonResponse({"ok": False, "fehler": str(exc)})

    return JsonResponse({"ok": True})


@login_required
def vorlage_oo_version(request, pk):
    vorlage = get_object_or_404(Briefvorlage, pk=pk)
    return JsonResponse({"version": vorlage.version})


def _vorlage_erstelle_wopi_token(vorlage: Briefvorlage) -> str:
    import secrets
    from django.utils import timezone
    from datetime import timedelta
    token = secrets.token_hex(32)
    Briefvorlage.objects.filter(pk=vorlage.pk).update(
        wopi_token=token,
        wopi_token_ablauf=timezone.now() + timedelta(hours=1),
    )
    return token


@csrf_exempt
def wopi_vorlage_files(request, pk):
    """WOPI CheckFileInfo fuer Briefvorlagen."""
    token = request.GET.get("access_token", "")
    vorlage = get_object_or_404(Briefvorlage, pk=pk)

    if not _vorlage_token_gueltig(vorlage, token):
        return HttpResponse("Unauthorized", status=401)

    from django.utils import timezone
    info = {
        "BaseFileName": f"Vorlage_{vorlage.titel}.docx",
        "Size": len(bytes(vorlage.inhalt)),
        "Version": str(vorlage.version),
        "OwnerId": str(vorlage.erstellt_von_id or ""),
        "UserId": str(request.user.pk) if request.user.is_authenticated else "",
        "UserFriendlyName": request.user.get_full_name() or request.user.username if request.user.is_authenticated else "Unbekannt",
        "UserCanWrite": True,
        "UserCanRename": False,
        "SupportsUpdate": True,
        "SupportsLocks": False,
        "LastModifiedTime": vorlage.erstellt_am.strftime("%Y-%m-%dT%H:%M:%SZ"),
    }
    return JsonResponse(info)


@csrf_exempt
def wopi_vorlage_contents(request, pk):
    """WOPI GetFile/PutFile fuer Briefvorlagen."""
    token = request.GET.get("access_token", "")
    vorlage = get_object_or_404(Briefvorlage, pk=pk)

    if not _vorlage_token_gueltig(vorlage, token):
        return HttpResponse("Unauthorized", status=401)

    if request.method == "GET":
        response = HttpResponse(bytes(vorlage.inhalt), content_type="application/octet-stream")
        response["Content-Disposition"] = f'attachment; filename="Vorlage_{vorlage.titel}.docx"'
        return response

    if request.method == "POST":
        inhalt = request.body
        if inhalt:
            Briefvorlage.objects.filter(pk=pk).update(
                inhalt=inhalt,
                version=vorlage.version + 1,
            )
        return JsonResponse({"LastModifiedTime": "2026-01-01T00:00:00Z"})

    return HttpResponse("Method Not Allowed", status=405)


def _vorlage_token_gueltig(vorlage: Briefvorlage, token: str) -> bool:
    from django.utils import timezone
    # Immer frisch aus DB lesen
    try:
        fresh = Briefvorlage.objects.values("wopi_token", "wopi_token_ablauf").get(pk=vorlage.pk)
    except Briefvorlage.DoesNotExist:
        return False
    t = fresh["wopi_token"]
    ablauf = fresh["wopi_token_ablauf"]
    return bool(t) and t == token and bool(ablauf) and timezone.now() < ablauf


# ---------------------------------------------------------------------------
# Views: Briefvorgaenge
# ---------------------------------------------------------------------------

@login_required
def brief_liste(request):
    briefe = Briefvorgang.objects.filter(erstellt_von=request.user).select_related("vorlage")
    return render(request, "korrespondenz/brief_liste.html", {"briefe": briefe})


@login_required
def brief_erstellen(request):
    vorlage_pk = request.GET.get("vorlage")
    if not vorlage_pk:
        standard = Briefvorlage.objects.filter(ist_aktiv=True, ist_standard=True).first()
        if standard:
            vorlage_pk = standard.pk

    initial = {}
    if vorlage_pk:
        initial["vorlage"] = vorlage_pk
        try:
            v = Briefvorlage.objects.get(pk=vorlage_pk, ist_aktiv=True)
            for f in ("absender_name", "absender_strasse", "absender_ort",
                      "absender_telefon", "absender_email", "ort", "grussformel"):
                val = getattr(v, f"default_{f}", None) or getattr(v, f, None)
                if val:
                    initial[f] = val
        except Briefvorlage.DoesNotExist:
            pass

    initial.setdefault("unterschrift_name", request.user.get_full_name() or request.user.username)

    form = BriefvorgangForm(request.POST or None, initial=initial)

    if request.method == "POST" and form.is_valid():
        vorgang = form.save(commit=False)
        vorgang.erstellt_von = request.user
        vorgang.save()

        try:
            platzhalter = _erstelle_platzhalter(vorgang)
            ausgefuellt = _fuelle_vorlage(bytes(vorgang.vorlage.inhalt), platzhalter)
            Briefvorgang.objects.filter(pk=vorgang.pk).update(inhalt=ausgefuellt)
        except Exception as exc:
            logger.error("Vorlage befuellen fehlgeschlagen fuer Brief %s: %s", vorgang.pk, exc)
            messages.warning(request, "Vorlage konnte nicht befuellt werden.")

        messages.success(request, "Brief erstellt.")
        return redirect("korrespondenz:brief_detail", pk=vorgang.pk)

    vorlagen_defaults = {}
    for v in Briefvorlage.objects.filter(ist_aktiv=True):
        vorlagen_defaults[v.pk] = {
            "absender_name":    v.default_absender_name,
            "absender_strasse": v.default_absender_strasse,
            "absender_ort":     v.default_absender_ort,
            "absender_telefon": v.default_absender_telefon,
            "absender_email":   v.default_absender_email,
            "ort":              v.default_ort,
            "grussformel":      v.default_grussformel,
        }

    return render(request, "korrespondenz/brief_erstellen.html", {
        "form": form,
        "vorlagen_defaults": vorlagen_defaults,
    })


@login_required
def brief_detail(request, pk):
    brief = get_object_or_404(Briefvorgang, pk=pk)
    editor_aktiv = bool(_get_oo_url()) and bool(brief.inhalt)
    return render(request, "korrespondenz/brief_detail.html", {
        "brief": brief,
        "collabora_aktiv": editor_aktiv,  # Template-Variable bleibt gleich
    })


@login_required
def brief_editor(request, pk):
    """Oeffnet einen Briefvorgang in OnlyOffice."""
    brief = get_object_or_404(Briefvorgang, pk=pk)
    oo_url = _get_oo_url()

    if not oo_url:
        messages.error(request, "OnlyOffice nicht konfiguriert (ONLYOFFICE_URL fehlt).")
        return redirect("korrespondenz:brief_detail", pk=pk)

    if not brief.inhalt:
        messages.error(request, "Kein DOCX-Inhalt vorhanden.")
        return redirect("korrespondenz:brief_detail", pk=pk)

    base_url = getattr(django_settings, "WOPI_BASE_URL",
                       getattr(django_settings, "VORGANGSWERK_BASE_URL",
                               request.build_absolute_uri("/").rstrip("/")))
    doc_key = f"korrespondenz-{brief.pk}-v{brief.version}"

    config = {
        "document": {
            "fileType": "docx",
            "key":      doc_key,
            "title":    f"Brief_{brief.datum}_{brief.betreff[:40]}.docx",
            "url":      f"{base_url}/korrespondenz/{brief.pk}/oo-download/",
        },
        "documentType": "word",
        "editorConfig": {
            "callbackUrl": f"{base_url}/korrespondenz/{brief.pk}/oo-callback/",
            "lang":        "de-DE",
            "region":      "de",
            "mode":        "edit",
            "user": {
                "id":   str(request.user.pk),
                "name": request.user.get_full_name() or request.user.username,
            },
            "customization": {"spellcheck": True},
        },
    }
    token = _oo_jwt(config)

    response = render(request, "korrespondenz/oo_editor.html", {
        "titel":       brief.betreff,
        "zurueck_url": f"/korrespondenz/{brief.pk}/",
        "oo_url":      oo_url,
        "oo_config":   config,
        "oo_jwt":      token,
        "brief":       brief,
        "ist_vorlage": False,
        "standard_platzhalter": _pfad_platzhalter(brief.vorlage.pfad_kuerzel),
        "pfad_platzhalter":     [],
    })
    response["Cache-Control"] = "no-store"
    return response


def brief_oo_download(request, pk):
    """Liefert die DOCX-Datei an OnlyOffice (JWT-gesichert via Authorization-Header)."""
    if not _oo_jwt_verify(request):
        return HttpResponse("Unauthorized", status=401)
    brief = get_object_or_404(Briefvorgang, pk=pk)
    return HttpResponse(
        bytes(brief.inhalt),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@csrf_exempt
def brief_oo_callback(request, pk):
    """OnlyOffice Speicher-Callback fuer Briefvorgaenge."""
    import json, urllib.request as urlreq, jwt as pyjwt
    if request.method != "POST":
        return JsonResponse({"error": 0})
    if not _oo_jwt_verify(request):
        return JsonResponse({"error": 1})
    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": 1})

    status = data.get("status")
    if status not in (2, 6):
        return JsonResponse({"error": 0})

    download_url = data.get("url", "")
    if not download_url:
        return JsonResponse({"error": 0})

    oo_pub = getattr(django_settings, "ONLYOFFICE_URL", "").rstrip("/")
    oo_int = getattr(django_settings, "ONLYOFFICE_INTERNAL_URL", "").rstrip("/")
    if oo_pub and oo_int and download_url.startswith(oo_pub):
        download_url = download_url.replace(oo_pub, oo_int, 1)

    try:
        req = urlreq.Request(download_url)
        secret = getattr(django_settings, "ONLYOFFICE_JWT_SECRET", "")
        if secret:
            dl_token = pyjwt.encode({"url": download_url}, secret, algorithm="HS256")
            req.add_header("Authorization", f"Bearer {dl_token}")
        with urlreq.urlopen(req, timeout=30) as resp:
            inhalt = resp.read()
    except Exception as exc:
        logger.error("OO-Callback Brief %s Download-Fehler: %s", pk, exc)
        return JsonResponse({"error": 1})

    brief = get_object_or_404(Briefvorgang, pk=pk)
    Briefvorgang.objects.filter(pk=pk).update(inhalt=inhalt, version=brief.version + 1)
    return JsonResponse({"error": 0})


@login_required
def brief_oo_forcesave(request, pk):
    """Loest Forcesave im OnlyOffice CommandService aus."""
    import json, urllib.request as urlreq
    if request.method != "POST":
        return JsonResponse({"ok": False})
    brief = get_object_or_404(Briefvorgang, pk=pk)
    doc_key = f"korrespondenz-{brief.pk}-v{brief.version}"

    oo_int = (getattr(django_settings, "ONLYOFFICE_INTERNAL_URL", "")
              or getattr(django_settings, "ONLYOFFICE_URL", "")).rstrip("/")
    if not oo_int:
        return JsonResponse({"ok": False, "fehler": "OnlyOffice nicht konfiguriert"})

    payload = {"c": "forcesave", "key": doc_key}
    headers = {"Content-Type": "application/json"}
    if getattr(django_settings, "ONLYOFFICE_JWT_SECRET", ""):
        headers["Authorization"] = f"Bearer {_oo_jwt(payload)}"

    try:
        req = urlreq.Request(
            f"{oo_int}/coauthoring/CommandService.ashx",
            data=json.dumps(payload).encode(),
            headers=headers,
            method="POST",
        )
        with urlreq.urlopen(req, timeout=10) as resp:
            antwort = json.loads(resp.read())
        if antwort.get("error", 1) not in (0, 4):
            return JsonResponse({"ok": False, "fehler": str(antwort)})
    except Exception as exc:
        logger.error("Brief %s Forcesave-Fehler: %s", pk, exc)
        return JsonResponse({"ok": False, "fehler": str(exc)})

    return JsonResponse({"ok": True})


@login_required
def brief_oo_version(request, pk):
    brief = get_object_or_404(Briefvorgang, pk=pk)
    return JsonResponse({"version": brief.version})


@csrf_exempt
def wopi_brief_files(request, pk):
    """WOPI CheckFileInfo fuer Briefvorgaenge."""
    token = request.GET.get("access_token", "")
    brief = get_object_or_404(Briefvorgang, pk=pk)

    if not brief.wopi_token or brief.wopi_token != token or not brief.wopi_token_gueltig():
        return HttpResponse("Unauthorized", status=401)

    info = {
        "BaseFileName": f"Brief_{brief.pk}.docx",
        "Size": len(bytes(brief.inhalt)) if brief.inhalt else 0,
        "Version": str(brief.version),
        "OwnerId": str(brief.erstellt_von_id or ""),
        "UserId": str(request.user.pk) if request.user.is_authenticated else "",
        "UserFriendlyName": request.user.get_full_name() or request.user.username if request.user.is_authenticated else "Unbekannt",
        "UserCanWrite": True,
        "UserCanRename": False,
        "SupportsUpdate": True,
        "SupportsLocks": False,
        "LastModifiedTime": brief.geaendert_am.strftime("%Y-%m-%dT%H:%M:%SZ"),
    }
    return JsonResponse(info)


@csrf_exempt
def wopi_brief_contents(request, pk):
    """WOPI GetFile/PutFile fuer Briefvorgaenge."""
    token = request.GET.get("access_token", "")
    brief = get_object_or_404(Briefvorgang, pk=pk)

    if not brief.wopi_token or brief.wopi_token != token or not brief.wopi_token_gueltig():
        return HttpResponse("Unauthorized", status=401)

    if request.method == "GET":
        response = HttpResponse(bytes(brief.inhalt), content_type="application/octet-stream")
        response["Content-Disposition"] = f'attachment; filename="Brief_{brief.pk}.docx"'
        return response

    if request.method == "POST":
        inhalt = request.body
        if inhalt:
            Briefvorgang.objects.filter(pk=pk).update(
                inhalt=inhalt,
                version=brief.version + 1,
            )
            brief.refresh_from_db(fields=["geaendert_am"])
        return JsonResponse({"LastModifiedTime": brief.geaendert_am.strftime("%Y-%m-%dT%H:%M:%SZ")})

    return HttpResponse("Method Not Allowed", status=405)


@login_required
def brief_download(request, pk):
    brief = get_object_or_404(Briefvorgang, pk=pk)
    if not brief.inhalt:
        messages.error(request, "Kein DOCX-Inhalt vorhanden.")
        return redirect("korrespondenz:brief_detail", pk=pk)
    response = HttpResponse(bytes(brief.inhalt),
                            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="Brief_{brief.pk}.docx"'
    return response


@login_required
def brief_pdf_download(request, pk):
    """Liefert das signierte PDF eines Briefvorgangs."""
    brief = get_object_or_404(Briefvorgang, pk=pk)
    if not brief.signiert_pdf:
        messages.error(request, "Kein signiertes PDF vorhanden.")
        return redirect("korrespondenz:brief_detail", pk=pk)
    dateiname = f"Brief_{brief.pk}_signiert.pdf"
    response = HttpResponse(bytes(brief.signiert_pdf), content_type="application/pdf")
    response["Content-Disposition"] = f'inline; filename="{dateiname}"'
    return response


@login_required
def brief_xoev_export(request, pk):
    """Exportiert den verknuepften Antrag als XÖV-kompatibles XML."""
    import xml.etree.ElementTree as ET
    from django.utils.timezone import localtime

    brief = get_object_or_404(Briefvorgang, pk=pk)
    sitzung = brief.sitzung

    NS = "http://xoev.de/vorgangswerk/1.0"
    ET.register_namespace("xoev", NS)

    def el(tag):
        return ET.Element(f"{{{NS}}}{tag}")

    def sub(parent, tag, text=""):
        e = ET.SubElement(parent, f"{{{NS}}}{tag}")
        e.text = str(text) if text else ""
        return e

    root = el("Vorgang")
    root.set("version", "1.0")

    # Metadaten
    meta = ET.SubElement(root, f"{{{NS}}}Metadaten")
    sub(meta, "Vorgangsnummer", sitzung.vorgangsnummer if sitzung else f"BRIEF-{brief.pk:05d}")
    sub(meta, "Verfahren", sitzung.pfad.name if sitzung else brief.vorlage.titel)
    sub(meta, "Verfahrenskuerzel", sitzung.pfad.kuerzel if sitzung else "")
    sub(meta, "Eingangsdatum", sitzung.abgeschlossen_am.strftime("%Y-%m-%d") if sitzung and sitzung.abgeschlossen_am else "")
    sub(meta, "ExportDatum", localtime().strftime("%Y-%m-%dT%H:%M:%S"))
    sub(meta, "Status", sitzung.get_status_display() if sitzung else brief.get_status_display())
    if brief.bearbeiter_signiert_von:
        sub(meta, "SigniertVon", brief.bearbeiter_signiert_von.get_full_name() or brief.bearbeiter_signiert_von.username)
        sub(meta, "SigniertAm", brief.bearbeiter_signiert_am.strftime("%Y-%m-%dT%H:%M:%S") if brief.bearbeiter_signiert_am else "")

    # Antragsteller
    antragsteller = el("Antragsteller")
    root.append(antragsteller)
    if sitzung:
        if sitzung.user:
            sub(antragsteller, "Name", sitzung.user.get_full_name() or sitzung.user.username)
            sub(antragsteller, "Benutzername", sitzung.user.username)
        elif sitzung.email_anonym:
            sub(antragsteller, "Email", sitzung.email_anonym)
    else:
        sub(antragsteller, "Name", brief.empfaenger_name)

    # Empfaenger (Behoerde)
    empf = ET.SubElement(root, f"{{{NS}}}Behoerde")
    sub(empf, "Name", brief.absender_name)
    sub(empf, "Strasse", brief.absender_strasse)
    sub(empf, "Ort", brief.absender_ort)

    # Antragsdaten aus Sitzung
    if sitzung:
        antragsdaten = ET.SubElement(root, f"{{{NS}}}Antragsdaten")
        daten = sitzung.gesammelte_daten or {}
        label_map = {}
        for schritt in sitzung.pfad.schritte.all():
            for feld in (schritt.felder_json or []):
                fid = feld.get("id") or feld.get("feldId", "")
                if fid:
                    label_map[fid] = feld.get("label", fid)
        fim_map = {}
        for schritt in sitzung.pfad.schritte.all():
            for feld in (schritt.felder_json or []):
                fid = feld.get("id") or feld.get("feldId", "")
                if fid and feld.get("fim_id"):
                    fim_map[fid] = feld["fim_id"]
        for schluessel, wert in daten.items():
            feld_el = ET.SubElement(antragsdaten, f"{{{NS}}}Feld")
            feld_el.set("bezeichnung", label_map.get(schluessel, schluessel))
            feld_el.set("feldId", schluessel)
            if schluessel in fim_map:
                feld_el.set("fimId", fim_map[schluessel])
            feld_el.text = str(wert) if wert else ""

    # Schreiben
    schreiben = ET.SubElement(root, f"{{{NS}}}Schreiben")
    sub(schreiben, "Betreff", brief.betreff)
    sub(schreiben, "Datum", str(brief.datum))
    sub(schreiben, "Unterzeichner", brief.unterschrift_name)
    sub(schreiben, "Funktion", brief.unterschrift_titel)

    xml_bytes = ET.tostring(root, encoding="unicode", xml_declaration=False)
    xml_out = f'<?xml version="1.0" encoding="UTF-8"?>\n{xml_bytes}'

    vgnr = (sitzung.vorgangsnummer if sitzung and sitzung.vorgangsnummer else f"Brief-{brief.pk}")
    dateiname = f"{vgnr}_xoev.xml".replace(" ", "_")
    response = HttpResponse(xml_out.encode("utf-8"), content_type="application/xml; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{dateiname}"'
    return response


@login_required
def brief_status_aendern(request, pk):
    brief = get_object_or_404(Briefvorgang, pk=pk)
    if request.method == "POST":
        neuer_status = request.POST.get("status", "")
        if neuer_status in dict(Briefvorgang.STATUS_CHOICES):
            brief.status = neuer_status
            brief.save(update_fields=["status"])
            messages.success(request, "Status geaendert.")
    return redirect("korrespondenz:brief_detail", pk=pk)


@login_required
def brief_loeschen(request, pk):
    brief = get_object_or_404(Briefvorgang, pk=pk)
    if request.method == "POST":
        brief.delete()
        messages.success(request, "Brief geloescht.")
        return redirect("korrespondenz:brief_liste")
    return redirect("korrespondenz:brief_detail", pk=pk)


# ---------------------------------------------------------------------------
# Bescheid-Generator (Brief aus Formular-Sitzung erstellen)
# ---------------------------------------------------------------------------

@login_required
def bescheid_suche(request):
    """Suche nach abgeschlossenen Formular-Sitzungen zum Erstellen eines Bescheids."""
    from django.db.models import Q
    import datetime
    from formulare.models import AntrSitzung, AntrPfad

    sitzungen = []
    vorlagen_fuer_sitzung = []
    vorlagen_passend = []
    vorlagen_allgemein = []
    q = request.GET.get("q", "").strip()
    gewaehlt_pk = request.GET.get("sitzung")
    gewaehlt = None

    if q:
        sitzungen = list(
            AntrSitzung.objects.filter(
                status=AntrSitzung.STATUS_ABGESCHLOSSEN,
            ).filter(
                Q(vorgangsnummer__icontains=q)
                | Q(pfad__name__icontains=q)
                | Q(pfad__kuerzel__icontains=q)
                | Q(user__first_name__icontains=q)
                | Q(user__last_name__icontains=q)
                | Q(email_anonym__icontains=q)
            ).select_related("pfad", "user").order_by("-abgeschlossen_am")[:50]
        )

    if gewaehlt_pk:
        gewaehlt = AntrSitzung.objects.filter(
            pk=gewaehlt_pk,
            status=AntrSitzung.STATUS_ABGESCHLOSSEN,
        ).select_related("pfad", "user").first()
        if gewaehlt:
            kuerzel = gewaehlt.pfad.kuerzel or ""
            alle_vorlagen = Briefvorlage.objects.filter(ist_aktiv=True).order_by("titel")
            vorlagen_passend   = [v for v in alle_vorlagen if kuerzel and v.pfad_kuerzel == kuerzel]
            vorlagen_allgemein = [v for v in alle_vorlagen if not v.pfad_kuerzel]
            vorlagen_fuer_sitzung = vorlagen_passend + vorlagen_allgemein

    if request.method == "POST":
        sitzung_pk = request.POST.get("sitzung_pk")
        vorlage_pk = request.POST.get("vorlage_pk")
        sitzung = get_object_or_404(AntrSitzung, pk=sitzung_pk, status=AntrSitzung.STATUS_ABGESCHLOSSEN)
        vorlage = get_object_or_404(Briefvorlage, pk=vorlage_pk, ist_aktiv=True)

        antrag_daten = {k: str(v) for k, v in sitzung.gesammelte_daten.items()}
        antrag_daten["vorgangsnummer"] = sitzung.vorgangsnummer or f"ANT-{sitzung.pk:05d}"
        antrag_daten["pfad_name"]      = sitzung.pfad.name
        antrag_daten["antrag_datum"]   = (
            sitzung.abgeschlossen_am.strftime("%d.%m.%Y") if sitzung.abgeschlossen_am else ""
        )
        if sitzung.user:
            antrag_daten["antragsteller"] = sitzung.user.get_full_name() or sitzung.user.username
        elif sitzung.email_anonym:
            antrag_daten["antragsteller"] = sitzung.email_anonym
        else:
            antrag_daten["antragsteller"] = "Anonym"

        from .models import Firmendaten as FD
        firma = FD.laden()
        vorgang = Briefvorgang(
            vorlage             = vorlage,
            sitzung             = sitzung,
            absender_name       = vorlage.default_absender_name    or firma.firmenname,
            absender_strasse    = vorlage.default_absender_strasse or firma.strasse,
            absender_ort        = vorlage.default_absender_ort     or firma.plz_ort,
            absender_telefon    = vorlage.default_absender_telefon or firma.telefon,
            absender_email      = vorlage.default_absender_email   or firma.email,
            empfaenger_name     = antrag_daten.get("antragsteller", ""),
            empfaenger_strasse  = antrag_daten.get("strasse", ""),
            empfaenger_plz_ort  = antrag_daten.get("plz_ort", ""),
            ort                 = vorlage.default_ort,
            datum               = datetime.date.today(),
            betreff             = f"Bescheid – {sitzung.pfad.name} ({antrag_daten['vorgangsnummer']})",
            anrede              = "Sehr geehrte Damen und Herren,",
            brieftext           = "",
            grussformel         = vorlage.default_grussformel or "Mit freundlichen Gruessen",
            unterschrift_name   = request.user.get_full_name() or request.user.username,
            unterschrift_titel  = "",
            erstellt_von        = request.user,
        )
        vorgang.save()

        try:
            platzhalter = _erstelle_platzhalter(vorgang)
            platzhalter.update(antrag_daten)
            ausgefuellt = _fuelle_vorlage(bytes(vorlage.inhalt), platzhalter)
            Briefvorgang.objects.filter(pk=vorgang.pk).update(inhalt=ausgefuellt)
        except Exception as exc:
            logger.error("Bescheid befuellen fehlgeschlagen fuer Brief %s: %s", vorgang.pk, exc)
            messages.warning(request, "Vorlage konnte nicht befuellt werden.")

        messages.success(request, "Bescheid erstellt.")
        return redirect("korrespondenz:brief_editor", pk=vorgang.pk)

    return render(request, "korrespondenz/bescheid_suche.html", {
        "q":                     q,
        "sitzungen":             sitzungen,
        "gewaehlt":              gewaehlt,
        "vorlagen_passend":      vorlagen_passend,
        "vorlagen_allgemein":    vorlagen_allgemein,
        "vorlagen_fuer_sitzung": vorlagen_fuer_sitzung,
    })


@login_required
def bescheid_platzhalter_json(request, sitzung_pk):
    """Liefert alle verfuegbaren Platzhalter einer Sitzung als JSON."""
    from formulare.models import AntrSitzung
    sitzung = get_object_or_404(AntrSitzung, pk=sitzung_pk, status=AntrSitzung.STATUS_ABGESCHLOSSEN)
    felder = {}
    for schritt in sitzung.pfad.schritte.all():
        for feld in schritt.felder():
            fid = feld.get("id", "")
            if fid:
                felder[fid] = {
                    "label": feld.get("label", fid),
                    "wert":  sitzung.gesammelte_daten.get(fid, ""),
                }
    return JsonResponse({"felder": felder})
